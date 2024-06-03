import pymupdf4llm

# input.pdf를 Markdown 텍스트로 변환
md_text = pymupdf4llm.to_markdown(r"gpt_servieces\input.pdf",write_images=True)

# Markdown 텍스트를 UTF-8 인코딩으로 저장
with open("gpt_servieces\input2.md", "w", encoding="utf-8") as md_file:
    md_file.write(md_text)

print(md_file)

import markdown
from markdown.treeprocessors import Treeprocessor
from markdown.extensions import Extension
from docx import Document
from docx.shared import Inches
import re
import os

class ImgExtractor(Treeprocessor):
    def run(self, root):
        self.md.images = []
        for image in root.findall(".//img"):
            self.md.images.append(image.attrib["src"])

class ImgExtExtension(Extension):
    def extendMarkdown(self, md):
        img_ext = ImgExtractor(md)
        md.treeprocessors.register(img_ext, 'imgext', 15)

def convert_md_to_docx(md_file, output_file):
    # Markdown 파일 읽기
    with open(md_file, 'r', encoding='utf-8') as file:
        text = file.read()

    # Markdown을 HTML로 변환하면서 이미지 추출
    md = markdown.Markdown(extensions=[ImgExtExtension()])
    html = md.convert(text)

    # Word 문서 생성
    doc = Document()

    # 텍스트와 이미지를 Word 문서에 추가
    for block in html.split("\n\n"):
        if re.match(r'^<img', block):
            # 이미지 처리
            img_path = re.findall(r'src="([^"]+)"', block)[0]
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(5.5))
        else:
            # 텍스트 처리
            doc.add_paragraph(block.replace("<p>", "").replace("</p>", "").replace("<strong>", "").replace("</strong>", "").replace("<em>", "").replace("</em>", "").replace("<h1>", "").replace("</h1>", "").replace("<h2>", "").replace("</h2>", "").replace("<h3>", "").replace("</h3>", ""))

    # Word 파일 저장
    doc.save(output_file)
    print(f"Converted {md_file} to {output_file}")

# 예제 사용
md_file = 'example.md'
output_file = 'output.docx'

convert_md_to_docx(md_file, output_file)
